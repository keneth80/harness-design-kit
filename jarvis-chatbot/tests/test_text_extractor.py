"""문서 추출기 테스트 (PDF/docx/평문)."""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from src.memory.text_extractor import (
    UnsupportedFileError,
    extract_text,
)


# ─── 평문 ──────────────────────────────────────────
def test_extract_txt():
    with tempfile.NamedTemporaryFile(
        suffix=".txt", delete=False, mode="w", encoding="utf-8"
    ) as f:
        f.write("안녕 hello\n둘째 줄")
        p = Path(f.name)
    assert "안녕" in extract_text(p)
    assert "둘째 줄" in extract_text(p)


def test_extract_md():
    with tempfile.NamedTemporaryFile(
        suffix=".md", delete=False, mode="w", encoding="utf-8"
    ) as f:
        f.write("# 제목\n본문")
        p = Path(f.name)
    out = extract_text(p)
    assert "제목" in out and "본문" in out


def test_extract_no_extension_treated_as_plain():
    with tempfile.NamedTemporaryFile(suffix="", delete=False) as f:
        f.write("plain bytes".encode("utf-8"))
        p = Path(f.name)
    assert "plain bytes" in extract_text(p)


# ─── docx ─────────────────────────────────────────
def test_extract_docx_paragraphs(tmp_path):
    from docx import Document

    p = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("첫 단락 한글입니다.")
    doc.add_paragraph("Second paragraph in English.")
    doc.add_paragraph("")  # 빈 줄은 무시됨
    doc.save(str(p))
    out = extract_text(p)
    assert "첫 단락" in out
    assert "Second paragraph" in out


def test_extract_docx_table(tmp_path):
    from docx import Document

    p = tmp_path / "test.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "이름"
    table.rows[0].cells[1].text = "역할"
    table.rows[1].cells[0].text = "kenneth"
    table.rows[1].cells[1].text = "admin"
    doc.save(str(p))
    out = extract_text(p)
    assert "이름 | 역할" in out
    assert "kenneth | admin" in out


# ─── PDF ──────────────────────────────────────────
# 최소한의 유효 PDF (1페이지, "Hello PDF" 텍스트)
_MIN_PDF = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT /F1 24 Tf 100 700 Td (Hello PDF) Tj ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f\x20
0000000009 00000 n\x20
0000000058 00000 n\x20
0000000109 00000 n\x20
0000000228 00000 n\x20
0000000316 00000 n\x20
trailer
<< /Size 6 /Root 1 0 R >>
startxref
389
%%EOF
"""


def test_extract_pdf(tmp_path):
    p = tmp_path / "test.pdf"
    p.write_bytes(_MIN_PDF)
    out = extract_text(p)
    assert "Hello PDF" in out


def test_extract_invalid_pdf_raises(tmp_path):
    p = tmp_path / "bad.pdf"
    p.write_bytes(b"not a real pdf")
    with pytest.raises(UnsupportedFileError):
        extract_text(p)


# ─── 미지원 ────────────────────────────────────────
def test_unsupported_extension_raises(tmp_path):
    p = tmp_path / "x.exe"
    p.write_bytes(b"binary")
    with pytest.raises(UnsupportedFileError):
        extract_text(p)


def test_hwp_explicitly_unsupported(tmp_path):
    p = tmp_path / "doc.hwp"
    p.write_bytes(b"hwp content")
    with pytest.raises(UnsupportedFileError):
        extract_text(p)


def test_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        extract_text(tmp_path / "nope.txt")


# ─── RagStore 통합 ────────────────────────────────
def test_ragstore_rejects_unsupported(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("TELEGRAM_CHAT_ID", "109494677")
    monkeypatch.setenv("TELEGRAM_BOT_TOKEN", "t")
    from src.core import config as cfg

    cfg.get_settings.cache_clear()
    from src.memory.rag_store import RagStore

    store = RagStore()
    bad = tmp_path / "x.zip"
    bad.write_bytes(b"PK\x03\x04")
    with pytest.raises(UnsupportedFileError):
        store.ingest_file_for(bad, owner_id="kenneth", is_admin=True)


def test_ragstore_rejects_empty_text(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("TELEGRAM_CHAT_ID", "109494677")
    monkeypatch.setenv("TELEGRAM_BOT_TOKEN", "t")
    from src.core import config as cfg

    cfg.get_settings.cache_clear()
    from src.memory.rag_store import RagStore

    store = RagStore()
    empty = tmp_path / "empty.txt"
    empty.write_text("   \n\n  ")  # 공백만
    with pytest.raises(UnsupportedFileError):
        store.ingest_file_for(empty, owner_id="kenneth", is_admin=True)
