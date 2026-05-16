"""파일 → 텍스트 추출.

지원:
  - 평문 (.txt .md .csv .json .py .yml .yaml .html .log .rst): 직접 읽기
  - .pdf: pypdf (페이지별 텍스트 추출)
  - .docx: python-docx (paragraph + table 텍스트)

미지원 확장자는 UnsupportedFileError. .doc(구버전 Word), .hwp는 미지원.
"""
from __future__ import annotations

from pathlib import Path

from src.core.logger import get_logger

_log = get_logger("memory.extractor")


PLAIN_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".tsv",
    ".json",
    ".py",
    ".yml",
    ".yaml",
    ".html",
    ".htm",
    ".log",
    ".rst",
    ".xml",
    ".ini",
    ".toml",
    ".sh",
    ".sql",
    ".js",
    ".ts",
}


class UnsupportedFileError(Exception):
    """확장자별 추출기가 없을 때 발생."""


def extract_text(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    if suffix == "" or suffix in PLAIN_EXTENSIONS:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise UnsupportedFileError(
        f"지원되지 않는 형식: {suffix}. 평문/.pdf/.docx만 가능."
    )


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise UnsupportedFileError(f"PDF 파싱 실패: {e}") from e
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            _log.warning(f"pdf page {i} extract failed in {path.name}: {e}")
            text = ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document

    try:
        doc = Document(str(path))
    except Exception as e:
        raise UnsupportedFileError(f"docx 파싱 실패: {e}") from e
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
    # 표(table) 셀 텍스트도 포함
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
